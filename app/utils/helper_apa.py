import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List
import re


class APAValidator:
    def __init__(self):
        self.issues = []

    def validate_document(self, doc_path: str) -> List[str]:

        doc = docx.Document(doc_path)
        self.issues = []

        self._check_font(doc)
        self._check_margins(doc)
        self._check_line_spacing(doc)
        self._check_document_structure(doc)

        self._check_title_page(doc)
        self._check_abstract(doc)
        self._check_keywords(doc)
        self._check_main_text(doc)
        self._check_references(doc)
        self._check_header(doc)

        return self.issues

    def _check_font(self, doc):
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name != 'Times New Roman':
                    self.issues.append(f"Font is not Times New Roman: '{run.text}'")
                if run.font.size and run.font.size.pt != 12:
                    self.issues.append(f"Font size is not 12pt: '{run.text}'")

    def _check_margins(self, doc):
        sections = doc.sections
        for section in sections:
            if (section.left_margin.inches != 1 or
                    section.right_margin.inches != 1 or
                    section.top_margin.inches != 1 or
                    section.bottom_margin.inches != 1):
                self.issues.append("Margins are not set to 1 inch on all sides")

    def _check_line_spacing(self, doc):
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                if paragraph.paragraph_format.line_spacing != 2:
                    self.issues.append(f"Text is not double-spaced: '{paragraph.text}'")

                space_after = paragraph.paragraph_format.space_after
                space_before = paragraph.paragraph_format.space_before

                if (space_after is not None and space_after > 0) or (space_before is not None and space_before > 0):
                    self.issues.append(f"Extra space found between paragraphs: '{paragraph.text}'")

    def _check_document_structure(self, doc):
        required_sections = ['Title Page', 'Abstract', 'Keywords', 'References']
        found_sections = []

        for paragraph in doc.paragraphs:
            for section in required_sections:
                if section.lower() in paragraph.text.lower():
                    found_sections.append(section)

        missing_sections = set(required_sections) - set(found_sections)
        if missing_sections:
            self.issues.append(f"Missing required sections: {', '.join(missing_sections)}")

    def _check_title_page(self, doc):
        first_page = doc.paragraphs[:10]
        found_title = False
        for para in first_page:
            if para.style.name == 'Title':
                found_title = True
                if not self._is_title_case(para.text):
                    self.issues.append("Title is not in title case")
                if not self._is_centered(para):
                    self.issues.append("Title is not centered")
                if not any(run.bold for run in para.runs):
                    self.issues.append("Title is not bolded")
                break

        if not found_title:
            self.issues.append("Title not found in upper half of first page")

        author_info_found = False
        for para in first_page[1:]:
            if para.text.strip():
                author_info_found = True
                if not self._is_centered(para):
                    self.issues.append("Author information is not centered")
                break

        if not author_info_found:
            self.issues.append("Author information not found")

        author_note_found = False
        for para in first_page:
            if 'Author Note' in para.text:
                author_note_found = True
                if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    self.issues.append("Author Note is not centered")
                if not any(run.bold for run in para.runs):
                    self.issues.append("Author Note heading is not bolded")
                break

        if not author_note_found:
            self.issues.append("Author Note not found")

    def _check_abstract(self, doc):
        abstract_found = False
        for paragraph in doc.paragraphs:
            if 'abstract' in paragraph.text.lower():
                abstract_found = True
                if not self._is_centered(paragraph):
                    self.issues.append("Abstract heading is not centered")
                if not any(run.bold for run in paragraph.runs):
                    self.issues.append("Abstract heading is not bolded")
                words = len(paragraph.text.split())
                if words > 250:
                    self.issues.append("Abstract exceeds 250 words")
                break

        if not abstract_found:
            self.issues.append("Abstract section not found")

    def _check_keywords(self, doc):
        keywords_found = False
        for paragraph in doc.paragraphs:
            if 'keywords' in paragraph.text.lower():
                keywords_found = True

                if not paragraph.text.lower().startswith("keywords:"):
                    self.issues.append("Keywords heading should begin with 'Keywords:'")
                if not any(run.font.italic for run in paragraph.runs):
                    self.issues.append("Keywords heading is not italicized")

                if paragraph.paragraph_format.left_indent != Inches(0.5):
                    self.issues.append("Keywords heading is not indented 0.5 inches")

                content_found = False
                for para in doc.paragraphs:
                    if 'keywords:' in para.text.lower():
                        keywords = para.text.split(":")[1].strip()
                        if not keywords.islower():
                            self.issues.append("Keywords should be listed in lowercase")
                        if ',' not in keywords:
                            self.issues.append("Keywords should be separated by commas")
                        content_found = True
                        break

                if not content_found:
                    self.issues.append("Keywords content not found below 'Keywords:'")

                break

        if not keywords_found:
            self.issues.append("Keywords section not found")

    def _check_main_text(self, doc):
        main_text_started = False
        for i, paragraph in enumerate(doc.paragraphs):
            if 'keywords' in paragraph.text.lower():
                main_text_started = True
            elif 'references' in paragraph.text.lower() and main_text_started:
                break
            elif main_text_started:
                if i == 0:
                    if not paragraph.text.strip():
                        self.issues.append("Main text should start on a new page after 'Keywords'.")

        title_found = False
        first_paragraph = doc.paragraphs[0]
        if first_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and first_paragraph.bold:
            title_found = True
        if not title_found:
            self.issues.append(
                "The title should be repeated in bold and centered at the top of the first page of the main text.")

        citation_pattern = r'\(([\w\s&]+, \d{4}(?:, .+)?(?:, p. \d{1,3})?)\)'
        for paragraph in doc.paragraphs:
            if re.search(citation_pattern, paragraph.text):
                citations = re.findall(citation_pattern, paragraph.text)
                for citation in citations:
                    authors = citation.split(",")[0].strip()
                    if '&' in authors and len(authors.split('&')) > 2:
                        self.issues.append(f"More than two authors in citation should be in the form of 'Smith et al.'")
                    if 'et al.' in citation and len(authors.split()) == 1:
                        self.issues.append(
                            f"Correct citation format for multiple authors should be '(Smith et al., 2020)'")
                        if "p." in citation:
                            if not re.search(r'\(.*p\. \d+\)', citation):
                                self.issues.append(
                                    f"Direct quotes should include page number, e.g., '(Smith, 2020, p. 15)'.")

        first_heading_checked = False
        for paragraph in doc.paragraphs:
            if 'abstract' in paragraph.text.lower() or 'references' in paragraph.text.lower():
                continue

            if paragraph.style.name == 'Heading 1':
                if not first_heading_checked:
                    if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER or not paragraph.bold:
                        self.issues.append(f"First Level 1 heading should be centered and bold: {paragraph.text}")
                    first_heading_checked = True
                else:
                    pass
            elif paragraph.style.name == 'Heading 2':
                if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT or not paragraph.bold:
                    self.issues.append(f"Level 2 heading should be flush left and bold: {paragraph.text}")
            elif paragraph.style.name == 'Heading 3':
                if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT or not paragraph.bold or not paragraph.italic:
                    self.issues.append(f"Level 3 heading should be flush left, bold, and italic: {paragraph.text}")
            elif paragraph.style.name == 'Heading 4':
                if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT or not paragraph.bold or paragraph.text[-1] != '.':
                    self.issues.append(
                        f"Level 4 heading should be flush left, bold, ending with a period: {paragraph.text}")
            elif paragraph.style.name == 'Heading 5':
                if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT or not paragraph.bold or not paragraph.italic or \
                        paragraph.text[-1] != '.':
                    self.issues.append(
                        f"Level 5 heading should be flush left, bold, italic, ending with a period: {paragraph.text}")

        for table in doc.tables:
            for row in table.rows:
                if row.cells[0].paragraphs[0].text.strip():
                    if row.cells[0].paragraphs[0].alignment != WD_ALIGN_PARAGRAPH.LEFT:
                        self.issues.append("Table title should be flush left above the table")
                if any(cell.paragraphs[0].runs[0].bold for cell in row.cells):
                    self.issues.append("Table heading should be in bold")
            for row in table.rows:
                for cell in row.cells:
                    if cell._element.xpath('.//w:vAlign') != []:
                        self.issues.append(f"Table should not have vertical borders: {row.text}")
            if table.rows[0].cells[0].text.strip()[:6].lower() != "table":
                self.issues.append("Tables should be numbered consecutively starting with 'Table 1'")

        for i, paragraph in enumerate(doc.paragraphs):
            if "figure" in paragraph.text.lower():
                figure_found = True
                if not re.search(r"Figure \d+", paragraph.text):
                    self.issues.append("Figures should be numbered sequentially, e.g., 'Figure 1'.")
                if not re.search(r"\b[a-zA-Z0-9\s]+$", paragraph.text):
                    self.issues.append(f"Figure caption should be brief and italicized: {paragraph.text}")

    def _check_references(self, doc):
        references_found = False
        check_references = False

        book_pattern = r'^[A-Za-z, ]+\.\s\(\d{4}\)\.\s[A-Za-z\s]+(?:\.\s)?[A-Za-z\s]+(?:\.\s)?[A-Za-z]+[\.]{1}$'  # Match books
        journal_pattern = r'^[A-Za-z, ]+\.\s\(\d{4}\)\.\s[A-Za-z\s]+(?:\.\s)?[A-Za-z\s]+(?:,|\s)?\d{1,2}\([0-9]+\)[,\s]\d{1,3}-\d{1,3}[\.]{1}$'  # Match journal articles
        website_pattern = r'^[A-Za-z, ]+\.\s\(\d{4},\s[A-Za-z]{3}\s\d{1,2}\)\.\s[A-Za-z\s]+(?:\.\s)?[A-Za-z\s]+(?:\.\s)?https?://[A-Za-z0-9./-]+$'  # Match websites
        doi_pattern = r'^[A-Za-z, ]+\.\s\(\d{4}\)\.\s[A-Za-z\s]+(?:\.\s)?[A-Za-z\s]+(?:,|\s)?\d{1,2}\([0-9]+\)[,\s]\d{1,3}-\d{1,3}\shttps://doi.org/[A-Za-z0-9/.-]+$'  # Match DOI format

        for paragraph in doc.paragraphs:
            if 'references' in paragraph.text.lower():
                references_found = True
                if not self._is_centered(paragraph):
                    self.issues.append("References title should be centered")
                if not any(run.font.bold for run in paragraph.runs):
                    self.issues.append("References title should be bold")
                check_references = True
                continue

            if check_references:
                if paragraph.text.strip():
                    if paragraph.paragraph_format.line_spacing != 2:
                        self.issues.append("References should be double-spaced")

                    text = paragraph.text.strip()
                    if not (re.match(book_pattern, text) or re.match(journal_pattern, text) or
                            re.match(website_pattern, text) or re.match(doi_pattern, text)):
                        self.issues.append(f"Reference format incorrect: '{text}'")

                    if paragraph.paragraph_format.first_line_indent != Inches(-0.5):
                        self.issues.append("References should have a hanging indent of 0.5 inches")

        if not references_found:
            self.issues.append("References section not found")

    def _check_header(self, doc):
        for section in doc.sections:
            header = section.header
            running_head_found = False
            page_number_found = False

            for paragraph in header.paragraphs:
                if paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                    running_head_found = True
                    if paragraph.text != paragraph.text.upper():
                        self.issues.append("Running head should be in all uppercase letters")

                if paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT and 'page' in paragraph.text.lower():
                    page_number_found = True
                    if not any(run.text.isdigit() for run in paragraph.runs):
                        self.issues.append("Page number is missing or not correct")

            if not running_head_found:
                self.issues.append(
                    "Running head is missing or not properly formatted (should be on the left side of the header)")

            if not page_number_found:
                self.issues.append(
                    "Page number is missing or not properly formatted (should be on the right side of the header)")

    def _is_title_case(self, text: str) -> bool:
        words = text.split()
        for word in words:
            if not word[0].isupper():
                return False
        return True

    def _is_centered(self, paragraph) -> bool:
        return paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER

# validator = APAValidator()
# issues = validator.validate_document(doc_path)
#
# if issues:
#     print("Issues found:")
#     for issue in issues:
#         print(f"- {issue}")
# else:
#     print("No issues found.")
